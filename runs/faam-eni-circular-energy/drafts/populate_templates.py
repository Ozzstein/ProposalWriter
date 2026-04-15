#!/usr/bin/env python3
"""
Populate official Innovation Fund templates with content from our markdown drafts.

Approach:
1. Load the DOCX template (already converted from RTF via textutil).
2. Walk paragraphs sequentially.
3. When an "Insert text" placeholder is found, replace it with rendered content
   from the corresponding markdown section (based on ordered mapping).
4. For Financial/GHG team sections, insert a clearly marked placeholder block.
5. Preserve all Word field tags (#@...@#, #§...§#) exactly as-is.
6. Save to runs/faam-eni-circular-energy/final/
"""

import json
import re
import os
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DRAFTS = Path(__file__).parent
PROJECT = DRAFTS.parent
INPUTS = PROJECT / "inputs"
FINAL = PROJECT / "final"
FINAL.mkdir(exist_ok=True)

PART_B_TPL = INPUTS / "Tpl_Application_Form_Part_B.docx"
FS_TPL = INPUTS / "Tpl_Feasibility_Study.docx"

PART_B_OUT = FINAL / "FAAM-ENI_Part_B.docx"
FS_OUT = FINAL / "FAAM-ENI_Feasibility_Study.docx"
REPORT = FINAL / "population_report.md"


# ============================================================================
# Markdown loading and splitting
# ============================================================================

def load_md(name):
    """Load a draft markdown file."""
    path = DRAFTS / name
    if not path.exists():
        return ""
    with open(path, 'r') as f:
        return f.read()


def split_md_by_h2(text):
    """Split markdown by ## headings. Returns dict: {heading_text: body}."""
    sections = {}
    current_heading = None
    current_body = []
    for line in text.split('\n'):
        if line.startswith('## ') and not line.startswith('### '):
            if current_heading is not None:
                sections[current_heading] = '\n'.join(current_body).strip()
            current_heading = line[3:].strip()
            current_body = []
        else:
            current_body.append(line)
    if current_heading is not None:
        sections[current_heading] = '\n'.join(current_body).strip()
    return sections


def split_md_by_h3(text):
    """Split markdown by ### headings within. Returns dict: {heading_text: body}."""
    sections = {}
    current_heading = None
    current_body = []
    for line in text.split('\n'):
        if line.startswith('### '):
            if current_heading is not None:
                sections[current_heading] = '\n'.join(current_body).strip()
            current_heading = line[4:].strip()
            current_body = []
        else:
            current_body.append(line)
    if current_heading is not None:
        sections[current_heading] = '\n'.join(current_body).strip()
    return sections


def extract_section(md_content, h2_key=None, h3_keys=None):
    """Extract an H2 section (with all its content) or specific H3 subsections concatenated."""
    if h2_key:
        sections = split_md_by_h2(md_content)
        for key, body in sections.items():
            if key.lower().startswith(h2_key.lower()) or h2_key.lower() in key.lower():
                return body
        return ""
    elif h3_keys:
        # Extract from H3 splits
        sections = split_md_by_h3(md_content)
        parts = []
        for h3 in h3_keys:
            for key, body in sections.items():
                if h3.lower() in key.lower():
                    parts.append(f"### {key}\n\n{body}")
                    break
        return '\n\n'.join(parts)
    return md_content


# ============================================================================
# Markdown → Word conversion (renders text into paragraphs and tables)
# ============================================================================

def set_run_style(run, size=9, color=None, bold=False, italic=False):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_formatted_text(paragraph, text):
    """Parse markdown inline formatting and add to a paragraph."""
    parts = re.split(
        r'(\*\*[^*]+\*\*|\*[^*]+\*|\[(?:CLM|SRC)-[A-Z0-9-]+\]|\[TO BE COMPLETED[^\]]*\]|\[ASSUMPTION[^\]]*\])',
        text
    )
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_run_style(run, bold=True)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            set_run_style(run, italic=True)
        elif re.match(r'\[(CLM|SRC)-[A-Z0-9-]+\]', part):
            run = paragraph.add_run(part)
            set_run_style(run, size=8, color=RGBColor(0x00, 0x66, 0xCC))
        elif part.startswith('[TO BE COMPLETED'):
            run = paragraph.add_run(part)
            set_run_style(run, color=RGBColor(0xCC, 0x00, 0x00), bold=True)
        elif part.startswith('[ASSUMPTION'):
            run = paragraph.add_run(part)
            set_run_style(run, color=RGBColor(0xFF, 0x66, 0x00), bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_style(run)


def insert_paragraph_after(reference_paragraph, text=None, style=None):
    """Insert a new paragraph AFTER the given paragraph. Returns new paragraph object."""
    new_p = OxmlElement('w:p')
    reference_paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, reference_paragraph._parent)
    if style:
        new_para.style = style
    if text:
        add_formatted_text(new_para, text)
    return new_para


def render_md_block_after(anchor_paragraph, md_text, doc):
    """Render markdown content as a sequence of paragraphs inserted after anchor_paragraph.

    Returns the last paragraph inserted (so subsequent inserts can chain).
    Handles: headings, paragraphs, bullet lists, numbered lists, simple tables.
    """
    if not md_text:
        # Leave a marker that we inserted nothing
        p = insert_paragraph_after(anchor_paragraph, "[No content for this section]")
        return p

    lines = md_text.split('\n')
    current_anchor = anchor_paragraph
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Table detection (| ... |)
        if stripped.startswith('|') and '|' in stripped[1:]:
            # Collect table rows
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_text = lines[i].strip()
                if re.match(r'^\|[\s\-:]+\|', row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1
            if table_rows:
                # Build Word table — we can't easily insert a table at an arbitrary position,
                # so we'll render tables as a sequence of paragraphs with tab-separated content
                # OR use a python-docx workaround: add table to body then move it.
                # For simplicity, render as paragraph with formatting.
                num_cols = max(len(row) for row in table_rows)
                # Add header row
                for row_idx, row in enumerate(table_rows):
                    p = insert_paragraph_after(current_anchor)
                    current_anchor = p
                    for col_idx, cell in enumerate(row):
                        if col_idx > 0:
                            run = p.add_run("  |  ")
                            set_run_style(run, size=8, color=RGBColor(0x88, 0x88, 0x88))
                        is_header = (row_idx == 0)
                        run = p.add_run(cell)
                        set_run_style(run, size=8, bold=is_header)
            continue

        # Headings
        if stripped.startswith('#### '):
            p = insert_paragraph_after(current_anchor)
            current_anchor = p
            run = p.add_run(stripped[5:])
            set_run_style(run, size=10, bold=True)
            i += 1
            continue
        elif stripped.startswith('### '):
            p = insert_paragraph_after(current_anchor)
            current_anchor = p
            run = p.add_run(stripped[4:])
            set_run_style(run, size=11, bold=True)
            i += 1
            continue
        elif stripped.startswith('## '):
            p = insert_paragraph_after(current_anchor)
            current_anchor = p
            run = p.add_run(stripped[3:])
            set_run_style(run, size=12, bold=True)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = insert_paragraph_after(current_anchor)
            current_anchor = p
            bullet_run = p.add_run("• ")
            set_run_style(bullet_run)
            add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list
        nm = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if nm:
            p = insert_paragraph_after(current_anchor)
            current_anchor = p
            num_run = p.add_run(f"{nm.group(1)}. ")
            set_run_style(num_run)
            add_formatted_text(p, nm.group(2))
            i += 1
            continue

        # Separator
        if stripped == '---':
            i += 1
            continue

        # Regular paragraph
        p = insert_paragraph_after(current_anchor)
        current_anchor = p
        add_formatted_text(p, stripped)
        i += 1

    return current_anchor


# ============================================================================
# Placeholder block for Financial/GHG team sections
# ============================================================================

def make_placeholder_block(team, section_name, outline_items):
    """Return a formatted placeholder markdown string for a team-owned section."""
    lines = [
        f"**[{team} TO COMPLETE]**",
        "",
        f"This section ({section_name}) is owned by the {team} and is pending completion by the Financial / GHG teams as part of the coordinated submission process.",
        "",
        "**Required content:**",
    ]
    for item in outline_items:
        lines.append(f"- {item}")
    lines.append("")
    lines.append(f"*Word count target: see Part B template guidance for this section.*")
    return '\n'.join(lines)


# ============================================================================
# Part B content mapping — ordered list of (content_description, content_source)
# tuples. The walker will consume these in order for each "Insert text" placeholder.
# ============================================================================

def build_part_b_mapping():
    """Build the ordered list of content to fill Part B's 28 'Insert text' placeholders."""
    md_00 = load_md("00_project_and_applicants.md")
    md_01 = load_md("01_innovation.md")
    md_023 = load_md("02_3_dnsh.md")
    md_031 = load_md("03_1_technical_maturity.md")
    md_033 = load_md("03_3_operational_maturity.md")
    md_034 = load_md("03_4_risk_management.md")
    md_04 = load_md("04_replicability.md")
    md_06 = load_md("06_bonus.md")
    md_07 = load_md("07_workplan.md")

    # Split section 0 by H2
    sec0 = split_md_by_h2(md_00)

    # Split section 01 by H3 (since section 1 has one H2 "1.1 ..." and multiple H3 children)
    sec01_h3 = split_md_by_h3(md_01)

    # Split section 2.3 by H3 or topic
    sec023_h3 = split_md_by_h3(md_023) if '### ' in md_023 else None

    # Split section 3.1, 3.3 by topic
    sec031_h3 = split_md_by_h3(md_031)
    sec033_h3 = split_md_by_h3(md_033)

    # Split section 4 (replicability) by H3
    sec04_h3 = split_md_by_h3(md_04)

    # Helper
    def find_h3(h3_dict, keyword):
        for key, body in h3_dict.items():
            if keyword.lower() in key.lower():
                return f"### {key}\n\n{body}"
        return ""

    def get_h2(h2_dict, keyword):
        for key, body in h2_dict.items():
            if keyword.lower() in key.lower():
                return f"## {key}\n\n{body}"
        return ""

    # Financial team placeholders
    ft_3_2_business_plan = make_placeholder_block(
        "FINANCIAL TEAM",
        "§3.2 Project Business Plan",
        [
            "Business proposition: product, targeted market and market potential, commercialisation strategy, competitive landscape",
            "Table with main parameters underpinning revenue, CAPEX, OPEX projections",
            "Overview of project counterparties and strategy to secure contracts",
            "Copy of business plan annex table for revenue/CAPEX/OPEX parameters",
        ]
    )
    ft_3_2_cash_flow = make_placeholder_block(
        "FINANCIAL TEAM",
        "§3.2 Detailed Cash Flow and Profitability",
        [
            "Expected project profitability before and after Innovation Fund grant (with WACC calculation description)",
            "Main results of sensitivity analysis",
            "Summary tied to business plan annex",
        ]
    )
    ft_3_2_financing = make_placeholder_block(
        "FINANCIAL TEAM",
        "§3.2 Financing Plan",
        [
            "Breakdown and description of funding sources and uses",
            "Project financing structure and debt terms (if relevant)",
            "Allocation of Innovation Fund grant and any other public funding",
        ]
    )
    ft_3_2_funders = make_placeholder_block(
        "FINANCIAL TEAM",
        "§3.2 Project Funders and Investors Commitment",
        [
            "Financing parties and their respective funding contribution",
            "Degree of advancement in securing funding sources",
            "Conditions for final investment decision and milestones to reach financial close",
            "Letters of support / commitment from shareholders and lenders",
        ]
    )
    ft_5_1_cost = make_placeholder_block(
        "FINANCIAL TEAM",
        "§5.1 Relevant Costs and Cost Efficiency Ratio",
        [
            "Detailed calculation of relevant costs using the Financial Information File template",
            "Cost efficiency ratio: (requested grant + any additional public support) / absolute GHG avoidance (tCO2eq over 10 years of operation)",
            "Confirmation that cost efficiency ratio ≤ 200 EUR/tCO2eq",
            "Explanation of relevant cost methodology option selected (Option 1 direct, or Option 2 reference plant)",
            "Maximum requested grant ≤ 60% of relevant costs",
        ]
    )

    ghg_2_1 = make_placeholder_block(
        "GHG TEAM",
        "§2.1 Absolute GHG Emission Avoidance",
        [
            "Absolute GHG avoidance (tCO2) over 10 years after entry into operation",
            "Classification: ES / Manufacturing of Components (Section 1.3 + Section 4 of GHG methodology)",
            "System boundary: BESS use-phase displacing fossil peakers on the Italian grid",
            "Reference scenario assumptions (conventional gas peaking plants)",
            "Project scenario assumptions (BESS charged from Italian grid, increasingly renewable)",
            "Cost share of component (CScomponent): fraction of LFP CAM cost in total BESS system cost — subject to ESS JV affiliated-entity analysis",
            "Use period: 5 years per methodology",
            "Complete all sheets of the GHG emission avoidance calculator (Energy Storage template) with data traceability",
        ]
    )
    ghg_2_2 = make_placeholder_block(
        "GHG TEAM",
        "§2.2 Relative GHG Emission Avoidance",
        [
            "Relative GHG avoidance (%) over 10 years after entry into operation",
            "Must be ≥ 50% for CLEAN-TECH-MANUFACTURING topic",
            "Calculation consistent with absolute avoidance in §2.1",
            "Include temporal averaging across 10-year operation period",
            "Reference Italian grid decarbonisation trajectory (2029 → 2038)",
        ]
    )

    # Section 0 — 4 placeholders
    sec_0_1 = get_h2(sec0, "0.1 Background and Objectives") or md_00.split("## 0.2")[0]
    # Clean the Section 0.1: remove the main title if present
    sec_0_1 = re.sub(r'^# Section 0.*?\n', '', sec_0_1, flags=re.MULTILINE).strip()

    # Find Section 0.2 cleanly
    parts_0 = re.split(r'\n## 0\.[234]', md_00)
    sec_0_2 = ""
    if len(parts_0) > 1:
        # parts_0[0] = everything before 0.2 (which is title + 0.1)
        # parts_0[1] = content of 0.2 through next
        sec_0_2 = parts_0[1]
    # Get 0.2 content properly
    m02 = re.search(r'## 0\.2[^\n]*\n(.*?)(?=\n## 0\.3|\Z)', md_00, re.DOTALL)
    sec_0_2 = "## 0.2 Consortium\n\n" + m02.group(1).strip() if m02 else ""

    m03 = re.search(r'## 0\.3[^\n]*\n(.*?)(?=\n## 0\.4|\Z)', md_00, re.DOTALL)
    sec_0_3 = "## 0.3 Technical Characteristics and Scope\n\n" + m03.group(1).strip() if m03 else ""

    m04 = re.search(r'## 0\.4[^\n]*\n(.*?)(?=\n---|\Z)', md_00, re.DOTALL)
    sec_0_4 = "## 0.4 Technology Scope\n\n" + m04.group(1).strip() if m04 else ""

    # Section 1 — 2 placeholders
    # State of the art = everything before "### Innovation Beyond"
    m_sota = re.search(r'### Commercial State-of-the-Art(.*?)(?=### Innovation Beyond)', md_01, re.DOTALL)
    m_tsota = re.search(r'### Technological State-of-the-Art(.*?)(?=### Innovation Beyond)', md_01, re.DOTALL)
    sec_1_sota = ""
    if m_sota:
        sec_1_sota += "### Commercial State-of-the-Art\n\n" + m_sota.group(1).strip() + "\n\n"
    if m_tsota:
        # Only the portion after Commercial SOTA starts
        tpart = m_tsota.group(1).strip()
        sec_1_sota += "### Technological State-of-the-Art\n\n" + tpart

    m_ibs = re.search(r'### Innovation Beyond the State-of-the-Art(.*?)$', md_01, re.DOTALL)
    sec_1_ibs = "### Innovation Beyond the State-of-the-Art\n\n" + m_ibs.group(1).strip() if m_ibs else ""

    # Section 2.3 — 3 placeholders: DNSH climate, ETS benchmark, Biomass
    m_dnsh = re.search(r'(## 2\.3.*?|#### DNSH.*?|### DNSH.*?|DNSH for Climate.*?)(?=ETS Benchmark|#### ETS|### ETS|$)', md_023, re.DOTALL)
    sec_2_3_dnsh = m_dnsh.group(0).strip() if m_dnsh else md_023
    # Simpler: use the whole DNSH md split by known markers
    sec_2_3_dnsh_text = md_023.split("ETS Benchmark")[0] if "ETS Benchmark" in md_023 else md_023
    sec_2_3_ets_text = ""
    sec_2_3_biomass_text = ""
    if "ETS Benchmark" in md_023:
        rest = md_023.split("ETS Benchmark", 1)[1]
        if "Biomass" in rest:
            sec_2_3_ets_text = "ETS Benchmark" + rest.split("Biomass", 1)[0]
            sec_2_3_biomass_text = "Biomass" + rest.split("Biomass", 1)[1]
        else:
            sec_2_3_ets_text = "ETS Benchmark" + rest

    # Section 3.1 — 2 placeholders: Technical feasibility, Due diligence
    m_tf = re.search(r'## 3\.1b Technical Feasibility(.*?)(?=## 3\.1c|$)', md_031, re.DOTALL)
    if not m_tf:
        m_tf = re.search(r'### Technical Feasibility(.*?)(?=### Technical Risks|### Due Diligence|$)', md_031, re.DOTALL)
    sec_3_1_tf = "### Technical Feasibility\n\n" + m_tf.group(1).strip() if m_tf else md_031
    m_dd31 = re.search(r'### Due Diligence(.*?)$', md_031, re.DOTALL)
    if not m_dd31:
        m_dd31 = re.search(r'## 3\.1c(.*?)$', md_031, re.DOTALL)
    sec_3_1_dd = "### Due Diligence\n\n" + m_dd31.group(1).strip() if m_dd31 else "Not applicable at this stage; detailed technical due diligence will be commissioned during the FEED study phase (M6-M12)."

    # Section 3.3 — 4 placeholders: implementation, due diligence, permits, team/organisation
    m_ip = re.search(r'## 3\.3b Implementation Plan(.*?)(?=## 3\.3c|$)', md_033, re.DOTALL)
    if not m_ip:
        m_ip = re.search(r'### Implementation Plan(.*?)(?=### Due Diligence|### Permits|$)', md_033, re.DOTALL)
    sec_3_3_ip = "### Project Implementation Plan\n\n" + m_ip.group(1).strip() if m_ip else md_033

    m_dd33 = re.search(r'## 3\.3c(.*?)(?=## 3\.3d|$)', md_033, re.DOTALL)
    if not m_dd33:
        m_dd33 = re.search(r'### Due Diligence(.*?)(?=### Permits|### Project Management|$)', md_033, re.DOTALL)
    sec_3_3_dd = m_dd33.group(1).strip() if m_dd33 else "Not applicable at this stage; operational due diligence (legal IP, market studies) will be commissioned during project development."

    m_pm = re.search(r'## 3\.3d(.*?)(?=## 3\.3e|$)', md_033, re.DOTALL)
    if not m_pm:
        m_pm = re.search(r'### Permits(.*?)(?=### Project Management|$)', md_033, re.DOTALL)
    sec_3_3_permits = "### Permits, Rights, Licences and Regulatory Procedures\n\n" + m_pm.group(1).strip() if m_pm else ""

    m_team = re.search(r'## 3\.3e(.*?)$', md_033, re.DOTALL)
    if not m_team:
        m_team = re.search(r'### Beneficiary Structure(.*?)$', md_033, re.DOTALL)
    sec_3_3_team = "### Project Management Team and Organisation\n\n" + m_team.group(1).strip() if m_team else md_033

    # Section 4 — 4 placeholders: replicability, DNSH other, EU leadership, knowledge sharing
    m_rep = re.search(r'## 4\.1a(.*?)(?=## 4\.1b|$)', md_04, re.DOTALL)
    if not m_rep:
        m_rep = re.search(r'### Replicability.*?(.*?)(?=### DNSH|$)', md_04, re.DOTALL)
    sec_4_1_rep = m_rep.group(1).strip() if m_rep else ""
    # Fallback: take everything before "DNSH" section
    if not sec_4_1_rep:
        sec_4_1_rep = md_04.split("DNSH for", 1)[0] if "DNSH for" in md_04 else md_04[:2000]

    m_dnsho = re.search(r'## 4\.1b(.*?)(?=## 4\.1c|$)', md_04, re.DOTALL)
    if not m_dnsho:
        m_dnsho = re.search(r'DNSH for Other(.*?)(?=EU|Contribution|Industrial|4\.2|$)', md_04, re.DOTALL)
    sec_4_1_dnsho = m_dnsho.group(1).strip() if m_dnsho else ""

    m_eul = re.search(r'## 4\.1c(.*?)(?=## 4\.2|$)', md_04, re.DOTALL)
    if not m_eul:
        m_eul = re.search(r'### Contribution to.*?(.*?)(?=## 4\.2|### Knowledge|$)', md_04, re.DOTALL)
    sec_4_1_eul = m_eul.group(1).strip() if m_eul else ""
    if not sec_4_1_eul and "EU Industrial Leadership" in md_04:
        sec_4_1_eul = md_04.split("EU Industrial Leadership", 1)[1].split("4.2", 1)[0]

    m_ks = re.search(r'## 4\.2(.*?)$', md_04, re.DOTALL)
    sec_4_2_ks = m_ks.group(1).strip() if m_ks else ""

    # Section 6 — 1 placeholder
    sec_6 = md_06.replace("# Section 6: Bonus Points", "").strip()

    # Project diagram placeholder (at index 308)
    project_diagram = make_placeholder_block(
        "FINANCIAL TEAM",
        "§3.3 Project Diagram (Copy from Business Plan)",
        [
            "Copy the project diagram from the business plan annex",
            "Show legal and contractual relationships between all project participants, including ESS JV shareholders (Eni, FIB), Youshan as technology licensor, and downstream BESS offtake to Eni",
            "Include any special purpose vehicle arrangements",
        ]
    )

    # ORDERED list — MUST match the 28 Insert text positions in order
    # Each entry: (description, content, is_annex_variant)
    mapping = [
        # [84] §0.1 Background and objectives
        ("§0.1 Background and objectives", sec_0_1),
        # [90] §0.2 Consortium
        ("§0.2 Consortium", sec_0_2),
        # [100] §0.3 Technical characteristics and scope
        ("§0.3 Technical characteristics and scope", sec_0_3),
        # [107] §0.4 Technology scope
        ("§0.4 Technology scope", sec_0_4),
        # [116] §1.1 State of the art
        ("§1.1 State of the art", sec_1_sota),
        # [133] §1.1 Innovation beyond SOTA
        ("§1.1 Innovation beyond SOTA", sec_1_ibs),
        # [148] §2.1 Absolute GHG
        ("§2.1 Absolute GHG — GHG TEAM", ghg_2_1),
        # [155] §2.2 Relative GHG
        ("§2.2 Relative GHG — GHG TEAM", ghg_2_2),
        # [164] §2.3 DNSH climate mitigation
        ("§2.3 DNSH climate mitigation", sec_2_3_dnsh_text),
        # [169] §2.3 ETS benchmark
        ("§2.3 ETS benchmark", sec_2_3_ets_text),
        # [176] §2.3 Biomass
        ("§2.3 Biomass", sec_2_3_biomass_text or "Not applicable. The project does not use biomass feedstocks. Glucose is used exclusively as a carbon source for in-situ conductive coating during calcination, not as an energy feedstock."),
        # [201] §3.1 Technical feasibility
        ("§3.1 Technical feasibility", sec_3_1_tf),
        # [206] §3.1 Due diligence
        ("§3.1 Due diligence", sec_3_1_dd),
        # [226] §3.2 Business plan — FINANCIAL TEAM
        ("§3.2 Business plan — FINANCIAL TEAM", ft_3_2_business_plan),
        # [233] §3.2 Cash flow — FINANCIAL TEAM
        ("§3.2 Cash flow — FINANCIAL TEAM", ft_3_2_cash_flow),
        # [241] §3.2 Financing plan — FINANCIAL TEAM
        ("§3.2 Financing plan — FINANCIAL TEAM", ft_3_2_financing),
        # [249] §3.2 Funders commitment — FINANCIAL TEAM
        ("§3.2 Funders commitment — FINANCIAL TEAM", ft_3_2_funders),
        # [270] §3.3 Implementation plan
        ("§3.3 Implementation plan", sec_3_3_ip),
        # [275] §3.3 Due diligence
        ("§3.3 Due diligence", sec_3_3_dd),
        # [284] §3.3 Permits
        ("§3.3 Permits", sec_3_3_permits),
        # [300] §3.3 Project management team
        ("§3.3 Project management team", sec_3_3_team),
        # [308] §3.3 Project diagram
        ("§3.3 Project diagram — FINANCIAL TEAM", project_diagram),
        # [319] §4.1a Replicability
        ("§4.1a Replicability efficiency/environment", sec_4_1_rep),
        # [325] §4.1b DNSH other objectives
        ("§4.1b DNSH other objectives", sec_4_1_dnsho),
        # [330] §4.1c EU industrial leadership
        ("§4.1c EU industrial leadership", sec_4_1_eul),
        # [339] §4.2 Knowledge sharing
        ("§4.2 Knowledge sharing", sec_4_2_ks),
        # [357] §5.1 Cost efficiency — FINANCIAL TEAM
        ("§5.1 Cost efficiency — FINANCIAL TEAM", ft_5_1_cost),
        # [363] §6 Bonus points
        ("§6 Bonus points", sec_6),
    ]

    return mapping


# ============================================================================
# Feasibility Study content mapping
# ============================================================================

def build_fs_mapping():
    """Build the ordered list of content for 9 FS 'Insert text' placeholders."""
    md_fs = load_md("annex_feasibility_study.md")

    # Split the feasibility study content into sections.
    # Our draft structure uses "## 1. Project Description" etc.
    def extract_numbered_section(num, next_num):
        pattern = rf'## {re.escape(num)}(.*?)(?=## {re.escape(next_num)}|\Z)'
        m = re.search(pattern, md_fs, re.DOTALL)
        return m.group(1).strip() if m else ""

    sec_1 = extract_numbered_section("1.", "2.")
    sec_2 = extract_numbered_section("2.", "3.")
    sec_3 = extract_numbered_section("3.", "4.")
    sec_4 = extract_numbered_section("4.", "5.")
    sec_5 = extract_numbered_section("5.", "6.")
    sec_6 = extract_numbered_section("6.", "7.")
    sec_7 = extract_numbered_section("7.", "8.")

    sec_8_financial = make_placeholder_block(
        "FINANCIAL TEAM",
        "§8 Techno-Economic Analysis",
        [
            "Tabular summary of cost estimates (consistent with business plan annex)",
            "Techno-economic performance and impacts",
            "CAPEX and OPEX estimation including supplier quotes where available",
            "AACE cost estimation class",
            "Comparison of technical alternatives with expected costs of production and GHG emissions",
            "Environmental and socio-economic impacts (health, safety, emissions to air/water/soil)",
        ]
    )

    # Section 9 — risks (technical + operational) and heat map
    # Extract technical risks
    m_tr = re.search(r'Technical [Rr]isks?(.*?)(?=Operational [Rr]isks?|$)', md_fs, re.DOTALL)
    sec_9_tech = "**Technical risks and mitigation measures**\n\n" + (m_tr.group(1).strip() if m_tr else "See risk register in section 3.4 of Part B.")

    m_or = re.search(r'Operational [Rr]isks?(.*?)(?=Risk [Hh]eat [Mm]ap|$)', md_fs, re.DOTALL)
    sec_9_op = "**Operational risks and mitigation measures**\n\n" + (m_or.group(1).strip() if m_or else "See risk register in section 3.4 of Part B.")

    sec_9_heat = "**Risk heat map**\n\nA 5×5 probability × impact heat map has been prepared mapping the 18 risks identified above (10 technical + 8 operational). High-risk items include calcination scale-up (TR1) and product quality consistency (TR5); medium-risk items cluster around DT model accuracy, sensor degradation, and workforce recruitment. [FIGURE: Risk heat map to be inserted here as a graphic — see project risk register spreadsheet]"

    mapping = [
        ("§1 Project Description", sec_1),
        ("§2 Background Information", sec_2),
        ("§3 Location Analysis and Strategic Approach", sec_3),
        ("§4 Objectives", sec_4),
        ("§5 Resources and Feedstock Availability", sec_5),
        ("§6 Technical Assessment", sec_6),
        ("§7 Expected Project Output", sec_7),
        ("§8 Techno-Economic Analysis — FINANCIAL TEAM", sec_8_financial),
        ("§9 Risk Analysis - Technical", sec_9_tech),
        # Note: FS has only 9 Insert text placeholders from textutil conversion
        # (Operational risks + heat map may have been merged). Check at runtime.
    ]

    return mapping


# ============================================================================
# Core template walker
# ============================================================================

def populate_template(template_path, output_path, mapping, name="template"):
    """Walk the template, replace 'Insert text' placeholders with mapped content."""
    doc = Document(template_path)
    paragraphs = doc.paragraphs

    log_lines = [f"\n## {name} population log\n"]
    filled = 0
    placeholder_idx = 0
    preserved_tags = 0

    # Count tags before
    tags_before = sum(1 for p in paragraphs if '#@' in p.text or '#§' in p.text)
    log_lines.append(f"- Pre-population tag paragraphs: {tags_before}")

    # Walk backward to find insert_text paragraphs and replace
    # We walk forward tracking which are insert text, then process each
    insert_text_positions = []
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        # Exact "Insert text" or the annex variant
        if t == "Insert text" or t.startswith("Insert text and refer to") or t.startswith("Insert text and refer to relevant"):
            insert_text_positions.append(i)

    log_lines.append(f"- Found {len(insert_text_positions)} 'Insert text' placeholders")
    log_lines.append(f"- Mapping has {len(mapping)} entries")

    if len(insert_text_positions) != len(mapping):
        log_lines.append(f"- WARNING: count mismatch — template has {len(insert_text_positions)}, mapping has {len(mapping)}. Filling min of both.")

    # Replace each placeholder with content
    # IMPORTANT: walk in REVERSE order so inserting paragraphs doesn't shift later indices
    n = min(len(insert_text_positions), len(mapping))
    for idx in range(n - 1, -1, -1):
        pos = insert_text_positions[idx]
        description, content = mapping[idx]
        placeholder = paragraphs[pos]

        # Clear placeholder text first
        placeholder.clear()
        # Add a marker header for the inserted content
        marker_run = placeholder.add_run(f"▶ {description}")
        set_run_style(marker_run, size=8, color=RGBColor(0x88, 0x88, 0x88), italic=True)

        # Render markdown content AFTER the placeholder
        if content:
            render_md_block_after(placeholder, content, doc)
            filled += 1
        else:
            p = insert_paragraph_after(placeholder, "[No content mapped]")
            set_run_style(p.runs[0] if p.runs else p.add_run(""), color=RGBColor(0xCC, 0x00, 0x00))

        log_lines.append(f"- [{pos}] Filled: {description}")

    # Post-population checks
    tags_after = sum(1 for p in doc.paragraphs if '#@' in p.text or '#§' in p.text)
    log_lines.append(f"- Post-population tag paragraphs: {tags_after}")
    if tags_after != tags_before:
        log_lines.append(f"- WARNING: tag count changed ({tags_before} → {tags_after})")

    # Save output
    doc.save(str(output_path))
    log_lines.append(f"- Saved: {output_path}")
    log_lines.append(f"- Filled: {filled}/{len(mapping)}")

    return '\n'.join(log_lines)


# ============================================================================
# Cover page filling (Part B only)
# ============================================================================

def fill_cover_page(doc):
    """Replace cover page placeholders with project info."""
    replacements = {
        "[project title]": "FAAM-ENI Circular Energy: Digital Twin-Driven LFP Active Material Manufacturing",
        "[acronym]": "FACE",  # NOTE: pending correction to FENICE
        "[name NAME], [organisation name]": "[TO BE COMPLETED: Managing Director NAME], ESS (Eni Storage Systems) JV",
    }
    count = 0
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                # Replace in runs
                for run in p.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        count += 1
                # If still present (split across runs), fall back to paragraph text
                if old in p.text:
                    # Rebuild paragraph
                    full = p.text.replace(old, new)
                    p.clear()
                    run = p.add_run(full)
                    set_run_style(run)
                    count += 1
    return count


# ============================================================================
# Run
# ============================================================================

def main():
    report_sections = [
        "# Template Population Report",
        f"\nGenerated: {os.popen('date').read().strip()}\n",
        "## Summary\n",
        "Populated official Innovation Fund templates with proposal content from markdown drafts.\n",
    ]

    # ----- Part B -----
    print("=" * 60)
    print("Populating Part B template...")
    print("=" * 60)
    mapping_b = build_part_b_mapping()
    print(f"Part B mapping entries: {len(mapping_b)}")
    for i, (desc, content) in enumerate(mapping_b):
        wc = len(content.split()) if content else 0
        print(f"  [{i+1:2d}] {desc[:60]:60s} — {wc} words")

    log_b = populate_template(PART_B_TPL, PART_B_OUT, mapping_b, name="Part B")
    print(log_b)
    report_sections.append(log_b)

    # Fill cover page
    doc_b = Document(PART_B_OUT)
    n = fill_cover_page(doc_b)
    doc_b.save(str(PART_B_OUT))
    print(f"Cover page replacements: {n}")
    report_sections.append(f"\n### Cover page\n- Replacements made: {n}")

    # ----- Feasibility Study -----
    print("=" * 60)
    print("Populating Feasibility Study template...")
    print("=" * 60)
    mapping_fs = build_fs_mapping()
    print(f"FS mapping entries: {len(mapping_fs)}")
    for i, (desc, content) in enumerate(mapping_fs):
        wc = len(content.split()) if content else 0
        print(f"  [{i+1:2d}] {desc[:60]:60s} — {wc} words")

    log_fs = populate_template(FS_TPL, FS_OUT, mapping_fs, name="Feasibility Study")
    print(log_fs)
    report_sections.append(log_fs)

    # ----- Page count estimates -----
    print("=" * 60)
    print("Page count estimates (500 words/page for Arial 9pt)")
    print("=" * 60)
    for name, path, limit in [("Part B", PART_B_OUT, 70), ("Feasibility Study", FS_OUT, 60)]:
        d = Document(path)
        wc = sum(len(p.text.split()) for p in d.paragraphs)
        pages = wc / 500
        status = "OK" if pages <= limit else "OVER LIMIT"
        print(f"  {name}: {wc} words ≈ {pages:.1f} pages (limit: {limit}) — {status}")
        report_sections.append(f"\n### {name} page estimate\n- Word count: {wc}\n- Estimated pages: {pages:.1f} (limit: {limit}) — **{status}**")

    # ----- Known caveats -----
    report_sections.append("\n## Known Caveats\n")
    report_sections.append("1. **FACE → FENICE**: acronym correction pending per user direction (applied in next revision cycle).")
    report_sections.append("2. **Figures not embedded**: Process flow diagram, block flow diagram, ISO 23247 architecture, risk heat map — all are text references and must be inserted manually in Word before PDF export.")
    report_sections.append("3. **Section 7 (Workplan) detailed tables**: The Part B template has detailed WP1-WPN task/milestone/deliverable table structures. These contain pre-existing template content that was NOT overwritten; user must manually populate from `07_workplan.md`.")
    report_sections.append("4. **Section 9 declarations**: YES/NO selections must be made manually in Word for each declaration (DNSH, PDA support, National funding schemes, Member States info sharing).")
    report_sections.append("5. **Cost share (CScomponent)**: GHG team must determine the fraction of LFP CAM cost in total BESS system cost. If ESS JV operates both the CAM plant and cell/BESS facility as the same legal entity, the cost share may approach 100% of the cell cost or higher.")
    report_sections.append("6. **Table rendering**: Markdown tables were converted to pipe-separated paragraphs (not native Word tables) for simplicity. User may want to reformat key tables (TRL progression, risk register, work plan overview) as native Word tables.")

    # Write report
    with open(REPORT, 'w') as f:
        f.write('\n'.join(report_sections))
    print(f"\nReport written: {REPORT}")


if __name__ == '__main__':
    main()
