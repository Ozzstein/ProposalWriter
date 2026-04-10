#!/usr/bin/env python3
"""Combine all proposal draft .md files into a single .docx with metadata comments."""

import json
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

DRAFTS_DIR = Path(__file__).parent

# Ordered list of draft files to include
SECTIONS = [
    "abstract.md",
    "00_project_and_applicants.md",
    "01_innovation.md",
    "02_3_dnsh.md",
    "03_1_technical_maturity.md",
    "03_3_operational_maturity.md",
    "03_4_risk_management.md",
    "04_replicability.md",
    "06_bonus.md",
    "07_workplan.md",
    "08_09_other_declarations.md",
]

ANNEX_SECTIONS = [
    "annex_feasibility_study.md",
]


def add_comment(paragraph, comment_text, author="ProposalWriter"):
    """Add a Word comment to a paragraph."""
    # Create comment reference
    comment_id = add_comment._next_id
    add_comment._next_id += 1

    # Add to comments part
    comments_part = paragraph.part.element.body.getparent()

    # Find or create comments element
    comments_el = comments_part.find(qn('w:comments'))
    if comments_el is None:
        # We need to use a different approach - add comments via the document
        # For simplicity, we'll add the metadata as a colored text block instead
        return None

    return comment_id

add_comment._next_id = 1


def add_metadata_block(doc, meta_path):
    """Add a colored metadata block from the _meta.json file."""
    if not meta_path.exists():
        return

    with open(meta_path, 'r') as f:
        meta = json.load(f)

    # Build comment text
    lines = []

    claims = meta.get('claim_ids', [])
    if claims:
        lines.append(f"CLAIMS REFERENCED: {', '.join(claims)}")

    assumptions = meta.get('assumptions_used', [])
    if assumptions:
        lines.append("ASSUMPTIONS:")
        for a in assumptions:
            lines.append(f"  - {a}")

    open_issues = meta.get('open_issues', [])
    if open_issues:
        lines.append("OPEN ISSUES:")
        for issue in open_issues:
            lines.append(f"  - {issue}")

    sources = meta.get('source_ids', [])
    if sources:
        lines.append(f"SOURCES CITED: {', '.join(sources[:20])}")
        if len(sources) > 20:
            lines.append(f"  ... and {len(sources) - 20} more")

    word_count = meta.get('word_count', meta.get('character_count', None))
    if word_count:
        lines.append(f"WORD COUNT: {word_count}")

    if not lines:
        return

    # Add a comment-style block
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Add border-like formatting via shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'FFF3CD')  # Light yellow background
    p.paragraph_format.element.get_or_add_pPr().append(shading)

    run = p.add_run("[REVIEWER NOTES] ")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x85, 0x63, 0x04)  # Dark amber

    for line in lines:
        run = p.add_run(line + "\n")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x85, 0x63, 0x04)


def md_to_docx_content(doc, md_text, is_annex=False):
    """Convert markdown text to docx paragraphs with basic formatting."""
    lines = md_text.split('\n')
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Headers
        if stripped.startswith('####'):
            p = doc.add_paragraph(stripped.lstrip('#').strip(), style='Heading 4')
            i += 1
            continue
        elif stripped.startswith('###'):
            p = doc.add_paragraph(stripped.lstrip('#').strip(), style='Heading 3')
            i += 1
            continue
        elif stripped.startswith('##'):
            p = doc.add_paragraph(stripped.lstrip('#').strip(), style='Heading 2')
            i += 1
            continue
        elif stripped.startswith('#'):
            p = doc.add_paragraph(stripped.lstrip('#').strip(), style='Heading 1')
            i += 1
            continue

        # Table detection
        if '|' in stripped and stripped.startswith('|'):
            # Collect table rows
            table_rows = []
            while i < len(lines) and '|' in lines[i].strip() and lines[i].strip().startswith('|'):
                row_text = lines[i].strip()
                # Skip separator rows (|---|---|)
                if re.match(r'^\|[\s\-:]+\|', row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1

            if table_rows:
                # Create table
                num_cols = max(len(row) for row in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols, style='Light Grid Accent 1')
                for r_idx, row in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < num_cols:
                            cell = table.cell(r_idx, c_idx)
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                paragraph.style.font.size = Pt(8)
                                for run in paragraph.runs:
                                    run.font.size = Pt(8)
                # Make first row bold
                if table_rows:
                    for cell in table.rows[0].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1


def add_formatted_text(paragraph, text):
    """Add text to a paragraph with basic markdown formatting (bold, italic)."""
    # Process bold and italic markers
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|\[(?:CLM|SRC|TO BE COMPLETED|ASSUMPTION)[^\]]*\])', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(9)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(9)
        elif part.startswith('[CLM-') or part.startswith('[SRC-'):
            run = paragraph.add_run(part)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)  # Blue for references
        elif part.startswith('[TO BE COMPLETED'):
            run = paragraph.add_run(part)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)  # Red for incomplete items
            run.bold = True
        elif part.startswith('[ASSUMPTION'):
            run = paragraph.add_run(part)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)  # Orange for assumptions
            run.bold = True
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(9)


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.59)  # 15mm
        section.bottom_margin = Inches(0.59)
        section.left_margin = Inches(0.59)
        section.right_margin = Inches(0.59)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FAAM-ENI Circular Energy')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x06, 0x16, 0x3A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Digital Twin-Driven LFP Active Material Manufacturing')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x06, 0x16, 0x3A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\nINNOVFUND-2025-NZT-CLEAN-TECH-MANUFACTURING\n')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Coordinator: FIB (FAAM)\nCo-investigator: Eni S.p.A.')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\nDRAFT — Generated {datetime.date.today().isoformat()}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\nNote: [CLM-XXX] = claim registry reference (blue)\n'
                     '[SRC-XXX] = evidence store reference (blue)\n'
                     '[TO BE COMPLETED: ...] = items pending (red)\n'
                     'Yellow blocks = reviewer metadata (claims, assumptions, open issues)')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # Part B sections
    p = doc.add_paragraph()
    run = p.add_run('PART B — TECHNICAL DESCRIPTION')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x06, 0x16, 0x3A)

    for md_file in SECTIONS:
        md_path = DRAFTS_DIR / md_file
        meta_path = DRAFTS_DIR / md_file.replace('.md', '_meta.json')

        if not md_path.exists():
            continue

        with open(md_path, 'r') as f:
            md_text = f.read()

        # Add section separator
        doc.add_page_break()

        # Add metadata block at top of section
        add_metadata_block(doc, meta_path)

        # Add section content
        md_to_docx_content(doc, md_text)

    # Annex sections
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run('ANNEX — FEASIBILITY STUDY')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x06, 0x16, 0x3A)

    for md_file in ANNEX_SECTIONS:
        md_path = DRAFTS_DIR / md_file
        meta_path = DRAFTS_DIR / md_file.replace('.md', '_meta.json')

        if not md_path.exists():
            continue

        with open(md_path, 'r') as f:
            md_text = f.read()

        doc.add_page_break()
        add_metadata_block(doc, meta_path)
        md_to_docx_content(doc, md_text, is_annex=True)

    # Save
    output_path = DRAFTS_DIR.parent / 'FAAM-ENI_Circular_Energy_DRAFT.docx'
    doc.save(str(output_path))
    print(f'Saved to: {output_path}')
    print(f'File size: {os.path.getsize(output_path) / 1024:.0f} KB')


if __name__ == '__main__':
    main()
