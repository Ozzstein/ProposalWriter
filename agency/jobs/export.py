"""export: assemble the proposal (markdown + DOCX) from the graph, with references."""
from __future__ import annotations

import re
from typing import Any

from agency.domain.graph import NodeType
from agency.domain.ids import CLAIM_REF_RE, SOURCE_REF_RE
from agency.domain.runs import JobKind
from agency.engine.plan import JobFailed, JobSpec, StageDef, StagePlan
from agency.engine.runtime import JobRuntime, RunContext
from agency.jobs import handler, stage


def plan_export(ctx: RunContext) -> StagePlan:
    return StagePlan("export", [
        JobSpec("assemble", "export.assemble", kind=JobKind.CODE),
        JobSpec("docx", "export.docx", kind=JobKind.CODE, deps=["assemble"], optional=True),
        JobSpec("finalize", "finalize_stage", kind=JobKind.GATE, deps=["assemble", "docx"], params={"gate": "submission"}),
    ])


stage(StageDef(name="export", state_key="export", planner=plan_export, requires_gate="submission",
               requires_stages=("review",),
               description="Assemble the full proposal as Markdown and DOCX with a reference list built from cited sources.",
               flags={"title": "override the document title"}))


def _reference(src) -> str:
    d = src.data
    bits = [d.get("authors") or "", f"({d.get('year')})" if d.get("year") else "", d.get("title", ""),
            d.get("doi") and f"doi:{d['doi']}" or d.get("url") or ""]
    return " ".join(b for b in bits if b).strip()


@handler("export.assemble")
async def assemble(rt: JobRuntime) -> dict[str, Any]:
    rt.ctx.materialize()
    spec = rt.ctx.callspec()
    project = rt.ws.get_project(rt.project_id)
    sections = rt.graph.sections()
    if not sections:
        raise JobFailed("no sections to export")
    title = rt.flags.get("title") or project.name
    cited_sources: list[str] = []
    cited_claims: set[str] = set()
    body_parts = []
    for s in sections:
        text = s.data.get("draft_text", "")
        cited_claims |= set(CLAIM_REF_RE.findall(text))
        for sid in SOURCE_REF_RE.findall(text):
            if sid not in cited_sources:
                cited_sources.append(sid)
        body_parts.append(text.strip())
    # sources cited indirectly through claims
    for cid in sorted(cited_claims):
        c = rt.graph.get(cid)
        for sid in (c.data.get("supported_by", []) if c else []):
            if sid not in cited_sources:
                cited_sources.append(sid)
    refs = []
    for i, sid in enumerate(cited_sources, 1):
        node = rt.graph.get(sid)
        if node:
            refs.append(f"{i}. [{sid}] {_reference(node)}")
    md = [f"# {title}", ""]
    if spec:
        md += [f"*{spec.funder} — {spec.title}*", ""]
    md += ["\n\n".join(body_parts), "", "## References", "", *refs]
    if cited_claims:
        md += ["", "## Claim register (internal)", ""]
        for cid in sorted(cited_claims):
            c = rt.graph.get(cid)
            if c:
                md.append(f"- {cid}: {c.data.get('text')} [{', '.join(c.data.get('supported_by', []))}]")
    text = "\n".join(md)
    out = rt.project_dir / "final"
    out.mkdir(exist_ok=True)
    (out / "proposal.md").write_text(text)
    key = rt.ws.blobs.put(text.encode(), suffix=".md")
    rt.graph.put_document("export_md", f"{title} (markdown)", text, created_by=rt.job.id, path=key,
                          sections=[s.data.get("section_id") for s in sections], references=len(refs))
    words = sum(len(p.split()) for p in body_parts)
    return {"path": str(out / "proposal.md"), "blob": key, "words": words, "references": len(refs),
            "summary": f"{len(sections)} sections, {words} words, {len(refs)} references"}


@handler("export.docx")
async def to_docx(rt: JobRuntime) -> dict[str, Any]:
    try:
        import docx
    except ImportError as e:  # pragma: no cover
        raise JobFailed(f"python-docx not installed: {e}")
    md_path = rt.project_dir / "final" / "proposal.md"
    text = md_path.read_text()
    doc = docx.Document()
    for raw in text.splitlines():
        line = raw.rstrip()
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            doc.add_heading(m.group(2), level=min(len(m.group(1)), 4))
        elif re.match(r"^\s*[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^\s*[-*]\s+", "", line), style="List Bullet")
        elif re.match(r"^\s*\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\s*\d+\.\s+", "", line), style="List Number")
        elif line.startswith("|"):
            doc.add_paragraph(line)
        elif line.strip():
            para = doc.add_paragraph()
            for part in re.split(r"(\*\*[^*]+\*\*)", line):
                run = para.add_run(part.strip("*") if part.startswith("**") else part)
                run.bold = part.startswith("**")
    out = rt.project_dir / "final" / "proposal.docx"
    doc.save(str(out))
    key = rt.ws.blobs.put_file(out)
    rt.graph.add(NodeType.DOCUMENT, {"kind": "export_docx", "title": out.name, "path": key}, created_by=rt.job.id)
    return {"path": str(out), "blob": key, "summary": f"DOCX written ({out.stat().st_size // 1024} KB)"}
